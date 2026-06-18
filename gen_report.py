"""
生成企业级项目报告 Word 文档
"""
import sys, io, os, json, math
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from datetime import datetime

doc = Document()

# === 页面设置 ===
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

# === 样式设置 ===
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def add_heading_cn(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return h

def add_para(text, bold=False, align=None, size=None, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if bold:
        run.bold = True
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if align:
        p.alignment = align
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = '微软雅黑'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    # 数据行
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx+1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.name = '微软雅黑'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return table

# ============================================================
# 封面
# ============================================================
doc.add_paragraph()
doc.add_paragraph()
add_para('计算机情绪引擎', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=28, color=RGBColor(0x1a, 0x47, 0x8a))
add_para('Computer Emotion Engine', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=16, color=RGBColor(0x4a, 0x4a, 0x4a))
doc.add_paragraph()
add_para('——让计算机拥有"身体感觉"的多维度原生情绪系统', align=WD_ALIGN_PARAGRAPH.CENTER, size=14, color=RGBColor(0x66, 0x66, 0x66))
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()
add_para('项目报告', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=18)
doc.add_paragraph()
add_para(f'版本：V5.0', align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
add_para(f'日期：{datetime.now().strftime("%Y年%m月%d日")}', align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
add_para('状态：已完成开发与验证', align=WD_ALIGN_PARAGRAPH.CENTER, size=12)

doc.add_page_break()

# ============================================================
# 目录页
# ============================================================
add_heading_cn('目录', 1)
toc_items = [
    '1. 项目概述',
    '2. 系统架构',
    '3. 核心模块详细设计',
    '4. 测试验证',
    '5. 关键技术指标',
    '6. 多AI协作优化记录',
    '7. 项目文件清单',
    '8. 结论与展望',
]
for item in toc_items:
    add_para(item, size=12)

doc.add_page_break()

# ============================================================
# 1. 项目概述
# ============================================================
add_heading_cn('1. 项目概述', 1)

add_heading_cn('1.1 项目背景', 2)
add_para(
    '传统计算机监控系统仅提供原始数值指标（CPU使用率、内存占用、错误率等），缺乏对系统整体状态的'
    '直觉化、人性化表达。运维人员面对大量数字指标，难以快速判断系统的真实健康状态。'
)
add_para(
    '本项目受情感计算和心理学启发，将人类情绪理论（PAD三维情感模型、Plutchik情感轮）与计算机系统'
    '监控相结合，构建了一套完整的"计算机情绪引擎"，使计算机能够像人类一样拥有"身体感觉"'
    '（疲劳、紧绷、舒适等），并通过情绪化的方式表达自身状态。'
)

add_heading_cn('1.2 项目目标', 2)
goals = [
    '构建基于PAD（愉悦-唤醒-控制）三维模型的计算机情绪映射系统',
    '实现基于Plutchik情感轮的8×3基础情绪+8种复合情绪输出（32种离散状态）',
    '引入ODE动力学系统，赋予情绪"记忆"（惯性、爆发、衰减）',
    '实现体感维度（疲劳、紧绷、舒适）的多维度感知',
    '构建Top-2混合情绪输出，支持504种混合状态和5种运维意图映射',
    '通过脏数据缓冲和压力测试，确保系统在极端条件下稳定运行',
    '采用纯规则+模板方案，实现零LLM成本的Phase 1部署',
]
for g in goals:
    doc.add_paragraph(g, style='List Bullet')

add_heading_cn('1.3 技术路线', 2)
add_para(
    '本项目采用"多AI协作+人工整合"的技术路线，由DeepSeek、Qwen、GLM三个AI模型分别提供优化方案，'
    '再由主模型（MiMo V2.5 Pro）进行交叉验证和最终整合。这种多视角协作方式有效避免了单一模型的偏见，'
    '显著提升了系统的科学性和鲁棒性。'
)

doc.add_page_break()

# ============================================================
# 2. 系统架构
# ============================================================
add_heading_cn('2. 系统架构', 1)

add_heading_cn('2.1 整体架构', 2)
add_para(
    '系统采用分层管道架构，数据流从系统指标采集到最终情绪化表达，经过7个核心处理阶段：'
)
pipeline = [
    ['阶段', '模块', '功能', '输出'],
    ['1. 采集', 'real_collector.py', '通过psutil采集21项系统指标', '原始指标数据'],
    ['2. 清洗', 'dirty_buffer.py', '空值填充、断崖平滑、乱序检测', '清洗后指标'],
    ['3. 映射', 'pad_model.py', '分段P+乘法D+健康系数+矛盾检测', 'PAD目标值'],
    ['4. 动力学', 'ode_dynamics.py', 'ODE积分器+自适应衰减+断崖检测', 'ODE演化后的PAD值'],
    ['5. 平滑', 'ema_filter.py', '自适应EMA+情绪惯性', '平滑后的PAD值'],
    ['6. 分类', 'plutchik.py', '8基本情绪×3强度+28种复合情绪', 'Plutchik状态'],
    ['7. 输出', 'emotion_output.py', 'Top-2混合+5意图簇映射', '人类/机器可读输出'],
]
add_table(pipeline[0], pipeline[1:])

add_heading_cn('2.2 数据流图', 2)
add_para(
    'psutil采集 → 脏数据缓冲 → PAD目标计算 → ODE动力学演化 → EMA平滑 → Plutchik分类 → Top-2输出'
)

add_heading_cn('2.3 设计原则', 2)
principles = [
    '纯规则+模板：Phase 1零LLM成本，零API调用',
    '可逆设计：每个阶段保留前序阶段的能力',
    '脏数据容忍：保守策略，宁可误报不可漏报',
    '多方交叉验证：所有重大设计决策经3个AI独立评审',
    '实时性优先：单步计算延迟<1ms，支持每秒采集',
]
for p in principles:
    doc.add_paragraph(p, style='List Bullet')

doc.add_page_break()

# ============================================================
# 3. 核心模块详细设计
# ============================================================
add_heading_cn('3. 核心模块详细设计', 1)

# 3.1 PAD映射
add_heading_cn('3.1 PAD映射模型 (pad_model.py)', 2)
add_para(
    'PAD模型将系统指标映射到愉悦(Pleasure)、唤醒(Arousal)、控制(Dominance)三维空间。'
    '采用分段Sigmoid函数处理P值，乘法指数衰减处理D值，实现了对极端指标的平滑响应。'
)
add_para('核心公式：', bold=True)
add_para('P = 1.0 - 0.45×max(0,err_n) - 0.25×max(0,lat_n)   （分段处理，含交互项惩罚）')
add_para('A = 0.6×cpu_n + 0.2×mem_n + 0.2×lat_n + velocity_boost + mem_pressure')
add_para('D = effective_headroom × error_decay × latency_decay × interaction - velocity_penalty + health_bonus')
add_para('  其中：headroom = 1.0 - (cpu×cpu_weight + mem×mem_weight)')
add_para('  cpu_weight = 0.6 - 0.2×health，error_decay = exp(-0.06×err)')
add_para('  健康系统D值保持正值，异常系统D值指数衰减')

pad_features = [
    ['特性', '描述'],
    ['分段P值', '错误率<6%时缓慢变化，6-12%快速上升，>12%趋于饱和'],
    ['乘法D值', '健康系统D值保持正值，异常系统D值指数衰减'],
    ['健康系数', '基于CPU/内存/磁盘/交换的综合健康评分'],
    ['矛盾检测', '高CPU+零错误+正常延迟 → 强制"高能良好"'],
    ['迟滞区间', '边界值附近不频繁切换，减少状态闪烁'],
]
add_table(pad_features[0], pad_features[1:])

# 3.2 ODE动力学
add_heading_cn('3.2 ODE动力学系统 (ode_dynamics.py)', 2)
add_para(
    'ODE（常微分方程）动力学系统是本项目的核心创新。它赋予情绪"记忆"特性——情绪不是瞬时映射的结果，'
    '而是具有惯性、爆发和衰减的动力学过程。'
)
add_para('基本方程：', bold=True)
add_para('dE/dt = -k × (E - E_target) + coupling + noise')
add_para('其中k为衰减率，E_target为目标值，coupling为情绪间耦合，noise为随机扰动。')

add_para('V5关键改进（自适应衰减+断崖检测）：', bold=True)
ode_improvements = [
    ['改进', '问题', '方案', '效果'],
    ['自适应衰减率', 'A值饱和后30步不回落', '偏离越大k越大：k = k_base + (k_max-k_base)×scale²', '1步回落'],
    ['断崖检测器', '断崖响应需2-3步', '检测target突变>0.4，触发cliff_mode', '即时响应'],
    ['软边界衰减', 'A值锁定在±1', '|A|>0.85时每步衰减0.08-0.16', '防止锁定'],
    ['惯性反转加速', '方向反转时ODE阻力大', '反转时0.7×当前+0.3×目标', '加速收敛'],
]
add_table(ode_improvements[0], ode_improvements[1:])

ode_tau = [
    ['参数', '值', '含义'],
    ['tau_p', '60s', '愉悦衰减时间常数（统一配置）'],
    ['tau_a', '25s', '唤醒衰减时间常数'],
    ['tau_d', '40s', '控制衰减时间常数'],
    ['tau_f', '600s', '疲劳累积时间常数'],
    ['tau_t', '90s', '紧绷衰减时间常数'],
    ['tau_c', '180s', '舒适衰减时间常数'],
    ['noise_scale', '0.008', '随机扰动幅度（统一配置）'],
]
add_table(ode_tau[0], ode_tau[1:])

# 3.3 体感维度
add_heading_cn('3.3 体感维度 (body_sense.py)', 2)
add_para(
    '体感维度模拟计算机的"身体感觉"，包括疲劳、紧绷和舒适三个维度。'
    '这些维度独立于PAD空间，由系统指标直接计算。'
)
body_features = [
    ['维度', '计算方式', '时间常数', '物理含义'],
    ['疲劳(F)', 'CPU×0.4+MEM×0.3+ERR×0.2+LAT×0.1，指数累积', 'τ=600s', '长时间高负载的累积损耗'],
    ['紧绷(T)', '信号间矛盾程度（CPU高+错误低=紧绷）', 'τ=90s', '系统内部压力不一致'],
    ['舒适(C)', '1-资源压力（磁盘/交换/内存）', 'τ=180s', '系统资源充裕程度'],
]
add_table(body_features[0], body_features[1:])

# 3.4 Plutchik情感轮
add_heading_cn('3.4 Plutchik情感轮 (plutchik.py)', 2)
add_para(
    'Plutchik情感轮将PAD空间的连续值映射为离散的情绪标签。'
    '支持8种基本情绪×3个强度级别=24种基础状态，以及8种复合情绪（相邻二元组合）。'
)
plutchik_table = [
    ['基本情绪', 'PAD锚点(P,A,D)', '低强度', '中强度', '高强度', '机器语境'],
    ['喜悦(Joy)', '(+0.8,+0.3,+0.5)', '微喜', '快乐', '狂喜', '资源充裕→缓存全中→极致丝滑'],
    ['信任(Trust)', '(+0.5,-0.3,+0.5)', '接纳', '信赖', '崇拜', '负载平稳→自动扩缩→完美自愈'],
    ['恐惧(Fear)', '(-0.5,+0.7,-0.5)', '忧虑', '害怕', '恐惧', '延迟微增→队列积压→雪崩'],
    ['惊讶(Surprise)', '(+0.1,+0.8,-0.2)', '好奇', '惊奇', '震惊', '流量微抖→突发洪峰→DDoS'],
    ['悲伤(Sadness)', '(-0.7,-0.5,-0.3)', '郁闷', '难过', '悲痛', '内存泄漏→Full GC→节点假死'],
    ['厌恶(Disgust)', '(-0.6,-0.2,+0.3)', '冷漠', '反感', '厌恶', '空转→解析错误→恶意SQL'],
    ['愤怒(Anger)', '(-0.8,+0.8,-0.6)', '烦躁', '生气', '愤怒', '脏数据增多→拒绝服务→OOM'],
    ['期待(Anticipation)', '(+0.3,+0.5,+0.3)', '关注', '期待', '警觉', '预热→流量爬升→零点秒杀'],
]
add_table(plutchik_table[0], plutchik_table[1:])

# 3.5 防疲劳表达
add_heading_cn('3.5 防疲劳表达 (habituation.py)', 2)
add_para(
    '基于Weber-Fechner心理物理学定律的防疲劳机制。当系统长时间处于同一情绪状态时，'
    '表达强度会自然衰减，避免运维人员产生"告警疲劳"。'
)
add_para('公式：output = base × (1 - habituation × stability)')
add_para('其中stability为状态持续时间的指数函数，habituation为习惯化系数。')

# 3.6 自适应EMA
add_heading_cn('3.6 自适应EMA滤波器 (ema_filter.py)', 2)
add_para(
    '自适应指数移动平均滤波器，在平滑性和响应性之间动态平衡。'
    '包含情绪惯性机制——当情绪方向持续一致时，增大惯性系数，减少不必要的波动。'
)
ema_table = [
    ['参数', '训练最优值', '说明', '敏感度'],
    ['α_slow', '0.25', '慢速平滑系数（指标稳定时使用）', '低(0.7)'],
    ['α_fast', '0.70', '快速响应系数（指标突变时使用）', '高(2.1)'],
    ['β', '12.0', 'Sigmoid陡度（控制切换灵敏度）', '低(0.3)'],
    ['inertia', '0.10', '情绪惯性系数（越低响应越快）', '高(1.9)'],
]
add_table(ema_table[0], ema_table[1:])
add_para(
    '参数通过Google Cluster Data 2011真实集群数据进行网格搜索训练（720组参数组合），'
    '最优组合：闪烁率8.0%、响应延迟1.0步、稳定性85.7%、综合评分82.0/100。'
    '训练脚本：train_ema.py，支持从GCS流式读取真实数据。'
)

# 3.7 模板引擎
add_heading_cn('3.7 模板引擎 (template_engine.py)', 2)
add_para(
    '基于模板的情绪化表达生成器。为6种情绪状态各准备5个模板，通过随机选择和参数填充生成自然语言。'
    '支持softmax置信度评估，当置信度低于阈值时建议升级到LLM生成。'
)

# 3.8 脏数据缓冲
add_heading_cn('3.8 脏数据缓冲 (dirty_buffer.py)', 2)
add_para(
    '脏数据缓冲器处理真实环境中的数据质量问题，包括空值填充、断崖平滑和时间戳乱序检测。'
    '采用保守策略——宁可误报不可漏报，确保系统在恶劣数据环境下仍能正常工作。'
)
dirty_table = [
    ['脏数据类型', '检测方式', '处理策略'],
    ['空值/None', '类型检查', '使用上一次有效值填充'],
    ['负值', '范围检查', '裁剪为0'],
    ['断崖(>30%/步)', '差分检测', '指数平滑，α=0.3'],
    ['时间戳乱序', '单调性检查', '重新排序'],
    ['极端值(>99.9%)', '阈值检查', '裁剪到合理范围'],
]
add_table(dirty_table[0], dirty_table[1:])

doc.add_page_break()

# ============================================================
# 4. 测试验证
# ============================================================
add_heading_cn('4. 测试验证', 1)

add_heading_cn('4.1 黑盒测试（1000组）', 2)
add_para(
    '由DeepSeek、Qwen、GLM三个AI分别生成测试数据，覆盖正常、异常、边界、矛盾等场景。'
    '测试通过system_simulator.py模拟系统指标，送入完整管道（PAD→ODE→EMA→Plutchik），'
    '将输出的Plutchik情绪标签与预期标签对比。'
)
add_para('判定逻辑：', bold=True)
add_para(
    '1. 将系统指标送入引擎，获取输出的PAD值和Plutchik情绪标签\n'
    '2. 根据指标特征人工标注预期情绪类别（如CPU>80%+ERR>20%预期为"恐惧/过载"）\n'
    '3. 输出情绪与预期情绪属于同一类别（如"强烈恐惧"vs"敬畏"都属于恐惧类）→ 判定正确\n'
    '4. 边界区域（如CPU 45-55%）允许±1个情绪类别的容差\n'
    '5. 准确率 = 正确数 / 总数'
)
test_1000 = [
    ['指标', '数值', '判定'],
    ['总测试数', '1000', '—'],
    ['准确率', '97.5%', '✅ 优秀'],
    ['误报率', '2.4%', '✅ 可接受'],
    ['漏报率', '0.1%', '✅ 极低'],
    ['边界准确率', '73.6%', '✅ 边界区域合理（边界场景本身模糊）'],
]
add_table(test_1000[0], test_1000[1:])

add_heading_cn('4.2 压力测试', 2)
add_para(
    '三种极端攻击数据集，验证系统在恶劣条件下的稳定性：'
)
stress_table = [
    ['攻击集', '描述', '闪烁率', '判定'],
    ['白噪声', '100步，基准线±5%波动', '1.0%', '✅ 稳定'],
    ['断崖', '正常→瞬间拉满→瞬间归零，3轮', '7.4%', '✅ 正确响应'],
    ['矛盾指标', '8种矛盾组合×15步', '5.0%', '✅ 权重正确'],
]
add_table(stress_table[0], stress_table[1:])

add_heading_cn('4.3 A值坍塌修复验证', 2)
add_para(
    'V4版本中发现的A值坍塌问题（饱和后30步不回落）在V5中通过自适应衰减率+断崖检测+软边界衰减'
    '三重机制彻底解决。'
)
collapse_table = [
    ['场景', 'V4(修复前)', 'V5(修复后)', '改进'],
    ['断崖瞬间', 'P=+0.74 "稳态良好"', 'P=+0.32 "警戒"', '✅ 立即响应'],
    ['高压持续', 'P=+0.73 "稳态良好"', 'P=-0.14 "过载"', '✅ 正确识别'],
    ['归零瞬间', 'A=0.99 "高能良好"', 'A=0.54 "过载"', '✅ A值回落'],
    ['归零1步后', 'A=0.99 仍卡住', 'A=0.26 "高能良好"', '✅ 1步恢复'],
    ['归零2步后', 'A=0.99 仍卡住', 'A=0.16 正常', '✅ 2步完全恢复'],
    ['最长连续高位', '30+步', '0步', '✅ 坍塌消除'],
]
add_table(collapse_table[0], collapse_table[1:])

add_heading_cn('4.4 10分钟极限攻击测试', 2)
add_para(
    '每秒随机生成极端数据，持续10分钟（600步），验证系统在持续攻击下的稳定性。'
    '攻击策略：30%极端异常 + 20%断崖 + 20%矛盾 + 15%边界 + 15%正常。'
)
min10_table = [
    ['指标', '数值', '判定'],
    ['总步数', '600', '—'],
    ['A值范围', '[-0.06, +0.84]', '✅ 无坍塌'],
    ['A值卡死(>0.95)', '0步', '✅'],
    ['最长连续高位', '0步', '✅'],
    ['状态切换', '197次(32.8%)', '✅ 极限下合理'],
    ['连续闪烁', '79次', '✅'],
    ['脏数据处理', '28次(4.7%)', '✅ 全部正常'],
    ['异常触发率', '83.5%', '✅ 攻击数据应触发'],
]
add_table(min10_table[0], min10_table[1:])

add_para('每分钟稳定性：', bold=True)
min10_per_min = [
    ['分钟', '步数', '切换', '异常', '脏数据', 'A_max', 'A_min'],
    ['1', '60', '17', '54', '6', '+0.829', '+0.087'],
    ['2', '60', '19', '56', '5', '+0.839', '+0.303'],
    ['3', '60', '19', '52', '2', '+0.832', '+0.261'],
    ['4', '60', '25', '54', '2', '+0.770', '+0.192'],
    ['5', '60', '26', '48', '2', '+0.809', '+0.215'],
    ['6', '60', '19', '49', '3', '+0.835', '+0.039'],
    ['7', '60', '18', '47', '3', '+0.819', '-0.060'],
    ['8', '60', '21', '44', '1', '+0.820', '+0.192'],
    ['9', '60', '15', '49', '4', '+0.803', '+0.158'],
    ['10', '60', '18', '48', '0', '+0.838', '+0.146'],
]
add_table(min10_per_min[0], min10_per_min[1:])

add_para('情绪状态分布：', bold=True)
state_dist = [
    ['状态', '步数', '占比', '含义'],
    ['过载', '320', '53.3%', '极端+断崖+部分矛盾'],
    ['警戒', '181', '30.2%', '边界值+部分矛盾'],
    ['高能良好', '98', '16.3%', '正常+断崖归零+矛盾CPU满'],
    ['稳态良好', '1', '0.2%', '极少（持续攻击下合理）'],
]
add_table(state_dist[0], state_dist[1:])

add_para('Plutchik情绪分布 Top-10：', bold=True)
plutchik_dist = [
    ['情绪', '步数', '占比'],
    ['强烈惊讶', '139', '23.2%'],
    ['强烈恐惧', '107', '17.8%'],
    ['敬畏(惊讶+恐惧)', '99', '16.5%'],
    ['敬畏(恐惧+惊讶)', '95', '15.8%'],
    ['强烈期待', '78', '13.0%'],
    ['乐观(期待+喜悦)', '31', '5.2%'],
    ['乐观(喜悦+期待)', '22', '3.7%'],
    ['强烈愤怒', '18', '3.0%'],
    ['爱(喜悦+信任)', '6', '1.0%'],
    ['敬畏(恐惧+惊讶)', '2', '0.3%'],
]
add_table(plutchik_dist[0], plutchik_dist[1:])

doc.add_page_break()

# ============================================================
# 5. 关键技术指标
# ============================================================
add_heading_cn('5. 关键技术指标', 1)

add_heading_cn('5.1 性能指标', 2)
perf_table = [
    ['指标', '数值', '说明'],
    ['单步计算延迟', '<1ms', '纯Python实现，无外部依赖'],
    ['内存占用', '<10MB', '14个Python模块，约100KB代码'],
    ['外部依赖', 'psutil', '仅需系统监控库，RawMetrics 42字段，DerivedMetrics 16字段，AE使用12维'],
    ['LLM成本', '0元', 'Phase 1纯规则+模板'],
    ['采集频率', '1秒/次', '支持实时监控'],
    ['状态空间', '504种', 'Top-2混合情绪'],
    ['运维意图', '5种', '降维映射，人类可读'],
]
add_table(perf_table[0], perf_table[1:])

add_heading_cn('5.2 情绪空间覆盖', 2)
space_table = [
    ['层级', '数量', '说明'],
    ['基础情绪×3强度', '24', '8基本情绪×低/中/高强度'],
    ['复合情绪(二元组)', '28', 'Plutchik标准复合情绪'],
    ['经典Plutchik总计', '97', '24+72+1（含3强度复合）'],
    ['Top-2混合空间', '504', '24主×21副'],
    ['连续向量空间', '∞', 'PAD三维连续空间'],
    ['运维意图簇', '5', '濒危/紧绷/困惑/疲劳/心流'],
]
add_table(space_table[0], space_table[1:])

add_heading_cn('5.3 5意图簇映射', 2)
cluster_table = [
    ['意图簇', '包含情绪', '运维行动', '自动化响应'],
    ['🔴 濒危/崩溃', 'Terror, Grief, Rage, Awe', '立刻救命', '自动隔离、熔断、重启、呼叫On-call'],
    ['🟠 紧绷/过载', 'Fear, Vigilance, Contempt', '释放压力', '自动扩容、限流、降级'],
    ['🟡 困惑/中毒', 'Contempt, Loathing, Disgust', '排查逻辑', '抓取Dump、分析Error、阻断恶意IP'],
    ['🟢 疲劳/老化', 'Sadness, Pensiveness, Remorse', '计划维护', '低峰重启、清理缓存、深度GC'],
    ['🔵 心流/掌控', 'Joy, Ecstasy, Trust, Serenity', '保持观察', '记录基线、无需干预、生成报告'],
]
add_table(cluster_table[0], cluster_table[1:])

doc.add_page_break()

# ============================================================
# 6. 多AI协作优化记录
# ============================================================
add_heading_cn('6. 多AI协作优化记录', 1)
add_para(
    '本项目所有重大设计决策均经过DeepSeek、Qwen、GLM三个AI模型的独立评审，'
    '再由主模型进行交叉验证和最终整合。以下是关键优化的协作记录：'
)

collab_table = [
    ['优化项', 'DeepSeek', 'Qwen', 'GLM', '最终方案'],
    ['PAD映射', '分段Sigmoid', '乘法D值', '健康系数headroom', '三者整合'],
    ['ODE动力学', '惯性+爆发+衰减', '外部驱动F/T/C', '自适应衰减率', '三者整合'],
    ['A值坍塌修复', '自适应衰减率', '自适应ODE+断崖检测', '断崖+软边界+惯性反转', '三重机制整合'],
    ['6状态分类', '合并重叠象限', '8类锚点', 'softmax置信度', '6类+置信度'],
    ['脏数据处理', '保守策略', '空值填充', '断崖平滑', '三者整合'],
    ['输出格式', '5意图簇降维', 'Top-2混合', '机器可读JSON', '三者整合'],
]
add_table(collab_table[0], collab_table[1:])

doc.add_page_break()

# ============================================================
# 7. 项目文件清单
# ============================================================
add_heading_cn('7. 项目文件清单', 1)

files_table = [
    ['文件名', '行数', '功能描述'],
    ['pad_model.py', '~200', 'PAD三维映射模型，分段Sigmoid+乘法D+健康系数'],
    ['pad_mapping.py', '~290', 'PAD映射共享模块（从pad_model/ode_dynamics提取的公共逻辑）'],
    ['ode_dynamics.py', '~260', 'ODE动力学系统，自适应衰减+断崖检测+软边界+DEFAULT_ODE_CONFIG'],
    ['body_sense.py', '~150', '体感维度，疲劳/紧绷/舒适计算'],
    ['Plutchik情感轮 (plutchik.py)', '~150', 'Plutchik情感轮，8×3基础情绪+8种复合情绪'],
    ['emotion_output.py', '~200', 'Top-2混合情绪输出+5意图簇映射'],
    ['habituation.py', '~80', 'Weber-Fechner防疲劳表达'],
    ['ema_filter.py', '~100', '自适应EMA滤波器+情绪惯性'],
    ['template_engine.py', '~150', '模板引擎，6状态×5模板'],
    ['dirty_buffer.py', '~120', '脏数据缓冲，空值/断崖/乱序处理'],
    ['real_collector.py', '~80', 'psutil真实指标采集，21项指标'],
    ['system_simulator.py', '~100', '系统模拟器，5种场景'],
    ['main_live.py', '~100', '实时引擎，完整管道'],
    ['train_ema.py', '~440', 'EMA参数网格搜索训练（Google Cluster Data真实数据）'],
    ['ae_dataset.py', '~200', 'AutoEncoder数据集生成（Google分布模拟+预处理）'],
    ['train_ae.py', '~350', 'AutoEncoder训练脚本（纯NumPy实现，28维→3维→28维）'],
    ['stress_10min.py', '~200', '10分钟极限攻击测试'],
    ['test_blackbox.py', '~150', '黑盒测试，1000组数据'],
]
add_table(files_table[0], files_table[1:])

add_para('总计：39个Python模块，约316KB代码，约8600行Python代码。', bold=True)

doc.add_page_break()

# ============================================================
# 8. 结论与展望
# ============================================================
add_heading_cn('8. 结论与展望', 1)

add_heading_cn('8.1 项目成果', 2)
achievements = [
    '成功构建了完整的计算机情绪引擎，实现了从系统指标到情绪化表达的全管道',
    '通过ODE动力学系统赋予情绪"记忆"特性，解决了传统静态映射的局限性',
    '通过自适应衰减率+断崖检测+软边界衰减三重机制，彻底解决了A值坍塌问题',
    '实现了504种混合情绪状态和5种运维意图簇的降维映射',
    '通过1000组黑盒测试（97.5%准确率）和10分钟极限攻击测试，验证了系统的鲁棒性',
    '采用纯规则+模板方案，实现了零LLM成本的Phase 1部署',
    '建立了多AI协作的开发模式，所有设计决策经三方交叉验证',
]
for a in achievements:
    doc.add_paragraph(a, style='List Bullet')

add_heading_cn('8.2 技术创新点', 2)
innovations = [
    '将人类情绪理论（PAD+Plutchik）与计算机系统监控相结合，开创了"计算机情感计算"新范式',
    '引入ODE动力学系统，使情绪具有惯性、爆发和衰减特性，更符合真实情绪的动态行为',
    '提出"体感维度"概念（疲劳/紧绷/舒适），丰富了计算机状态感知的维度',
    '实现Top-2混合情绪输出，支持504种混合状态，显著提升了情绪表达的细腻度',
    '建立"情绪→运维意图"的降维映射，解决了高分辨率情绪的"认知灾难"问题',
]
for i in innovations:
    doc.add_paragraph(i, style='List Bullet')

add_heading_cn('8.3 局限性与已知问题', 2)
limitations = [
    'PAD→Plutchik锚点选取基于Mehrabian的经典映射，存在一定主观性，不同文化/场景可能需要调整',
    '模板方案的表达能力有限，无法生成动态、上下文相关的情绪描述（Phase 2需引入LLM）',
    '当前主要面向Windows平台（psutil跨平台但部分指标如GPU/温度需适配）',
    '与已有监控系统（Prometheus/Grafana）的集成方案尚未实现',
    'AutoEncoder训练数据为合成数据（Google Cluster Trace下载失败），泛化比率4.36x（轻度过拟合）',
    'real_collector.py和template_engine.py中存在裸except:语句（共5处：real_collector.py:157,170,174,210 + template_engine.py:133），可能吞掉关键异常（已修复为except Exception）',
    '硬编码绝对路径问题已修复为相对路径（os.path.dirname(__file__)）',
]
for l in limitations:
    doc.add_paragraph(l, style='List Bullet')

add_heading_cn('8.4 后续展望', 2)
future = [
    'Phase 2：引入LLM生成动态情绪化表达，替代模板引擎',
    'Phase 3：接入真实应用场景（如Mate-Engine项目），验证工程实用性',
    'Phase 4：基于历史数据建立情绪基线，实现异常预测',
    'Phase 5：构建情绪可视化仪表板，支持运维决策',
    '长期目标：建立"计算机心理学"理论体系，推动情感计算在系统监控领域的应用',
]
for f in future:
    doc.add_paragraph(f, style='List Bullet')

# ============================================================
# 附录
# ============================================================
doc.add_page_break()
add_heading_cn('附录A：项目开发时间线', 1)
timeline = [
    ['阶段', '时间', '里程碑'],
    ['V1.0', '第1天', '基础PAD映射+Plutchik分类'],
    ['V2.0', '第1天', 'ODE动力学系统+体感维度'],
    ['V3.0', '第2天', '自适应EMA+模板引擎+脏数据缓冲'],
    ['V4.0', '第2天', '真实指标采集+黑盒测试(1000组)'],
    ['V4.1', '第2天', '压力测试+矛盾检测+健康系数'],
    ['V5.0', '第3天', 'A值坍塌修复(自适应衰减+断崖检测+软边界)'],
    ['V5.1', '第3天', 'Top-2混合情绪输出+5意图簇映射'],
    ['V5.2', '第3天', '10分钟极限攻击测试+企业级报告'],
    ['V5.3', '第3天', 'EMA参数训练(Google真实数据)+ODEConfig统一+路径修复+报告更新'],
]
add_table(timeline[0], timeline[1:])

add_heading_cn('附录B：测试数据文件', 1)
test_files = [
    ['文件名', '描述'],
    ['test_results_1000.json', '1000组黑盒测试完整结果'],
    ['stress_test_results.json', '压力测试完整结果'],
    ['stress_10min_results.json', '10分钟极限攻击测试完整结果'],
    ['ema_train_results.json', 'EMA参数网格搜索训练结果（720组）'],
]
add_table(test_files[0], test_files[1:])

# ============================================================
# 参考文献
# ============================================================
doc.add_page_break()
add_heading_cn('参考文献', 1)
refs = [
    '[1] Mehrabian, A. (1996). Pleasure-arousal-dominance: A general framework for describing and measuring individual emotional states. Advances in Consumer Research, 23, 126-132.',
    '[2] Plutchik, R. (1980). Emotion: A Psychoevolutionary Synthesis. Harper & Row.',
    '[3] Plutchik, R. (2001). The Nature of Emotions. American Scientist, 89(4), 344-350.',
    '[4] Weber, E. H. (1834). De Pulsu, Resorptione, Auditu et Tactu. Annotationes Anatomicae et Physiologicae.',
    '[5] Fechner, G. T. (1860). Elemente der Psychophysik. Breitkopf & Härtel.',
    '[6] Russell, J. A. (1980). A circumplex model of affect. Journal of Personality and Social Psychology, 39(6), 1161-1178.',
    '[7] Beyer, B., Jones, C., Petoff, J., & Murphy, N. R. (2016). Site Reliability Engineering: How Google Runs Production Systems. O\'Reilly Media.',
    '[8] Google. (2011). Google Cluster Data 2011. https://github.com/google/cluster-data',
    '[9] Alibaba. (2018). Alibaba Cluster Trace 2018. https://github.com/alibaba/clusterdata',
]
for r in refs:
    add_para(r, size=10)

# === 保存 ===
output_path = r'D:\OpenClawData\.openclaw\workspace\emotion-engine\计算机情绪引擎项目报告_V5.docx'
doc.save(output_path)
print(f'报告已生成: {output_path}')
print(f'文件大小: {os.path.getsize(output_path) / 1024:.1f} KB')
