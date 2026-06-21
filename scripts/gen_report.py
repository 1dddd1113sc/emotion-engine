"""
生成企业级项目报告 Word 文档 — V6.3
"""
import os
import sys, io, os, json, math
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from datetime import datetime

doc = Document()

# === 页面设置 ===
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

# === 样式设置 ===
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def add_heading_cn(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return h

def add_para(text, bold=False, align=None, size=None, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if bold:
        run.bold = True
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if align:
        p.alignment = align
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = '微软雅黑'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx+1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.name = '微软雅黑'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return table

# ============================================================
# 封面
# ============================================================
doc.add_paragraph()
doc.add_paragraph()
add_para('计算机情绪引擎', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=28, color=RGBColor(0x1a, 0x47, 0x8a))
add_para('Computer Emotion Engine', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=16, color=RGBColor(0x4a, 0x4a, 0x4a))
doc.add_paragraph()
add_para('——让计算机拥有"身体感觉"的多维度原生情绪系统', align=WD_ALIGN_PARAGRAPH.CENTER, size=14, color=RGBColor(0x66, 0x66, 0x66))
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()
add_para('项目报告', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=18)
doc.add_paragraph()
add_para(f'版本：V6.3', align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
add_para(f'日期：{datetime.now().strftime("%Y年%m月%d日")}', align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
add_para('状态：真机压力测试通过，7天连续采集已部署', align=WD_ALIGN_PARAGRAPH.CENTER, size=12)

doc.add_page_break()

# ============================================================
# 目录页
# ============================================================
add_heading_cn('目录', 1)
toc_items = [
    '1. 项目概述',
    '2. 系统架构（V6 四层管线）',
    '3. 核心模块详细设计',
    '4. V6.2 修复清单',
    '5. V6.3 新增：ODE-Kalman 融合管线',
    '6. V6.3 新增：真机压力测试',
    '7. V6.3 新增：7 天连续采集',
    '8. 测试验证',
    '9. 关键技术指标',
    '10. 多AI协作优化记录',
    '11. 项目文件清单',
    '12. 结论与展望',
]
for item in toc_items:
    add_para(item, size=12)

doc.add_page_break()

# ============================================================
# 1. 项目概述
# ============================================================
add_heading_cn('1. 项目概述', 1)

add_heading_cn('1.1 项目背景', 2)
add_para(
    '传统计算机监控系统仅提供原始数值指标（CPU使用率、内存占用、错误率等），缺乏对系统整体状态的'
    '直觉化、人性化表达。运维人员面对大量数字指标，难以快速判断系统的真实健康状态。'
)
add_para(
    '本项目受情感计算和心理学启发，将人类情绪理论（PAD三维情感模型、Plutchik情感轮）与计算机系统'
    '监控相结合，构建了一套完整的"计算机情绪引擎"，使计算机能够像人类一样拥有"身体感觉"'
    '（疲劳、紧绷、舒适等），并通过情绪化的方式表达自身状态。'
)

add_heading_cn('1.2 项目目标', 2)
goals = [
    '构建基于PAD（愉悦-唤醒-控制）三维模型的计算机情绪映射系统',
    '实现基于Plutchik情感轮的8×3基础情绪+8种复合情绪输出（32种离散状态）',
    '引入ODE动力学系统，赋予情绪"记忆"（惯性、爆发、衰减）',
    '实现体感维度（疲劳、紧绷、舒适）的多维度感知',
    '构建Top-2混合情绪输出，支持504种混合状态和5种运维意图映射',
    '通过脏数据缓冲和压力测试，确保系统在极端条件下稳定运行',
    '采用纯规则+模板方案，实现零LLM成本的Phase 1部署',
]
for g in goals:
    doc.add_paragraph(g, style='List Bullet')

add_heading_cn('1.3 技术路线', 2)
add_para(
    '本项目采用"多AI协作+人工整合"的技术路线，由DeepSeek、Qwen、GLM三个AI模型分别提供优化方案，'
    '再由主模型进行交叉验证和最终整合。V6 引入上下文感知架构，实现了"同一指标在不同上下文下产生不同情绪"'
    '的核心创新。'
)

doc.add_page_break()

# ============================================================
# 2. 系统架构
# ============================================================
add_heading_cn('2. 系统架构（V6 四层管线）', 1)

add_heading_cn('2.1 整体架构', 2)
add_para(
    'V6 采用四层浓缩架构，将 87 个原始指标逐步提炼为 7 维情感状态：'
)

add_para('四层管线数据流：', bold=True)
pipeline = [
    ['层级', '模块', '输入', '输出', '核心思想'],
    ['L1 采集层', 'real_collector.py', '系统API (psutil/WMI/nvidia-smi)', '87指标 (58原始+29派生)', '五层感官全覆盖'],
    ['L2 体感层', 'body_sense.py', '87指标', '3维度 (疲劳/紧绑/舒适)', 'EMA累积，有记忆'],
    ['L3 语义层', 'semantic_signals.py + context_pad.py', '原始指标 + 体感', 'PAD (P/A/D/V)', '上下文感知：同一load不同情绪'],
    ['L4 动力层', 'ode_dynamics.py + ema_filter.py + plutchik.py', 'PAD目标值', '7维情感 + Plutchik 24态', 'ODE惯性/爆发/衰减'],
]
add_table(pipeline[0], pipeline[1:])

add_heading_cn('2.2 上下文感知核心创新', 2)
add_para(
    'V6 的核心创新是"上下文感知 PAD 映射"。同一个 load 信号，在不同上下文下产生完全不同的情绪：'
)
ctx_table = [
    ['上下文', '条件', 'load↑ 的情绪', '含义'],
    ['clean', '系统健康，错误少', 'P↑ A↑ D↑ (自信/高效)', '忙碌但健康 = 高效运转'],
    ['degraded', '性能下降，无严重错误', 'P→ A↑ D↓ (紧张)', '勉强支撑 = 紧张'],
    ['err', '有严重错误', 'P↓ A↑ D↓ (焦虑/愤怒)', '过载 = 恐慌'],
]
add_table(ctx_table[0], ctx_table[1:])

add_heading_cn('2.3 设计原则', 2)
principles = [
    '纯规则+模板：Phase 1零LLM成本，零API调用',
    '上下文感知：同一指标在不同语境下产生不同情绪',
    '体感累积：BodySense 用 EMA 追踪疲劳/紧绑/舒适，有"记忆"',
    '脏数据容忍：保守策略，宁可误报不可漏报',
    '多方交叉验证：所有重大设计决策经3个AI独立评审',
    '实时性优先：单步计算延迟<1ms，支持每秒采集',
]
for p in principles:
    doc.add_paragraph(p, style='List Bullet')

doc.add_page_break()

# ============================================================
# 3. 核心模块详细设计
# ============================================================
add_heading_cn('3. 核心模块详细设计', 1)

# 3.1 五层感官采集
add_heading_cn('3.1 五层感官采集 (real_collector.py)', 2)
add_para(
    '采集器覆盖 5 个感官层级，共 58 个原始指标 + 29 个派生指标。'
)
layer_table = [
    ['层级', '名称', '指标数', '数据源', '情绪映射'],
    ['L1', '计算与记忆层', '21', 'psutil', 'Fatigue (疲劳)'],
    ['L2', '吞吐与排队层', '6', 'psutil', 'Stress (压力)'],
    ['L3', '传导与IO层', '17', 'psutil + WMI', 'Stress + Flow'],
    ['L4', '业务表现层', '8', 'Prometheus (需注入)', 'Flow / Confusion'],
    ['L5', '物理硬件层', '6', 'nvidia-smi + LHM', '终极 Fatigue'],
]
add_table(layer_table[0], layer_table[1:])

# 3.2 体感维度
add_heading_cn('3.2 体感维度 (body_sense.py)', 2)
add_para(
    '体感管理器将 87 个指标浓缩为 3 个维度，使用 EMA 累积追踪，具有"记忆"特性。'
)
body_table = [
    ['维度', '追踪器', '时间常数', '信号来源', '物理含义'],
    ['疲劳(F)', 'FatigueTracker', 'τ=600s (10分钟半衰期)', 'L1+L5: CPU过劳/降频/温度/GPU', '长时间高负载的累积消耗'],
    ['紧绑(T)', 'TensionTracker', '窗口=10步', 'L2+L3: 连接积压/IO延迟/中断', '系统内部压力矛盾程度'],
    ['舒适(C)', 'ComfortTracker', '即时计算', 'L3+L5: 磁盘空间/Swap/温度/IO', '系统资源充裕程度'],
]
add_table(body_table[0], body_table[1:])

# 3.3 语义信号
add_heading_cn('3.3 语义信号层 (semantic_signals.py)', 2)
add_para(
    '从原始指标提取 4 个正交语义信号，每个信号输出 [0, 1]，语义独立不重叠。'
)
sig_table = [
    ['信号', '含义', '计算方式'],
    ['error', '错误严重度', '分段升级：低错误缓慢上升，高错误急剧上升'],
    ['load', '综合负载', '0.55×CPU + 0.30×MEM + 0.15×Swap，高负载非线性放大'],
    ['latency', '延迟压力', 'tanh 归一化，只取正值，延迟飙升加重'],
    ['health', '健康余量', '0.45×err_health + 0.25×lat_health + 0.15×disk + 0.15×swap'],
]
add_table(sig_table[0], sig_table[1:])

# 3.4 上下文 PAD
add_heading_cn('3.4 上下文感知 PAD (context_pad.py)', 2)
add_para(
    'PAD 组合器接收语义信号和体感状态，根据上下文标签（clean/degraded/err）调整映射权重。'
    'V6.2 修复：compose_pad() 新增 body 参数，优先使用 BodySenseManager 的 EMA 累积值。'
)

# 3.5 ODE动力学
add_heading_cn('3.5 ODE动力学系统 (ode_dynamics.py)', 2)
add_para(
    'ODE（常微分方程）动力学系统赋予情绪"记忆"特性——情绪不是瞬时映射的结果，'
    '而是具有惯性、爆发和衰减的动力学过程。'
)
add_para('基本方程：', bold=True)
add_para('dE/dt = -k × (E - E_target) + coupling + noise')
add_para('V6.2：compute_target() 委托给 pad_mapping.compute_pad_raw()，消除 ~80 行重复代码。')

ode_tau = [
    ['参数', '值', '含义'],
    ['tau_p', '60s', '愉悦衰减时间常数'],
    ['tau_a', '25s', '唤醒衰减时间常数'],
    ['tau_d', '40s', '控制衰减时间常数'],
    ['tau_f', '600s', '疲劳累积时间常数'],
    ['tau_t', '90s', '紧绑衰减时间常数'],
    ['tau_c', '180s', '舒适衰减时间常数'],
    ['noise_scale', '0.008', '随机扰动幅度'],
]
add_table(ode_tau[0], ode_tau[1:])

# 3.6 EMA
add_heading_cn('3.6 自适应EMA滤波器 (ema_filter.py)', 2)
add_para(
    '自适应指数移动平均滤波器，在平滑性和响应性之间动态平衡。'
    'V6.2 修复：默认参数改为 V6 训练最优值 (α_slow=0.35, α_fast=0.60)。'
)
ema_table = [
    ['参数', 'V6 默认值', '说明'],
    ['α_slow', '0.35', '慢速平滑系数（指标稳定时使用）'],
    ['α_fast', '0.60', '快速响应系数（指标突变时使用）'],
    ['β', '12.0', 'Sigmoid 陡度'],
    ['inertia', '0.20', '情绪惯性系数'],
]
add_table(ema_table[0], ema_table[1:])

# 3.7 防闪烁
add_heading_cn('3.7 防闪烁象限控制器 (quadrant_stabilizer.py)', 2)
add_para(
    '四层防抖机制：维度死区 → 象限滞回 → 惯性窗口 → 震荡抑制。'
    '支持上下文自适应：clean 时放大死区+高惯性，err 时缩小死区+低惯性。'
    'V6.2 修复：新增 oscillation_suppress 属性，震荡检测不再崩溃。'
)

# 3.8 Plutchik
add_heading_cn('3.8 Plutchik情感轮 (plutchik.py)', 2)
add_para(
    '将 PAD 空间的连续值映射为离散情绪标签。8种基本情绪×3个强度=24种基础状态，'
    '加8种复合情绪（相邻二元组合）。'
)

# 3.9 其他模块
add_heading_cn('3.9 辅助模块', 2)
aux_table = [
    ['模块', '功能', 'V6.2 变更'],
    ['template_engine.py', '模板引擎，6状态×5模板', 'except 裸 catch → 具体异常类型；OutputThrottler 改用真实时间'],
    ['pad_model.py', 'PAD 状态、6 象限分类', 'docstring 版本更正为 V4.1'],
    ['pad_mapping.py', 'PAD 映射共享模块', 'docstring 版本更正为 V6.0'],
    ['dirty_buffer.py', '脏数据缓冲', '无变更'],
    ['habituation.py', 'Weber-Fechner 防疲劳表达', '无变更'],
    ['l4_metrics.py', 'L4 Prometheus 采集', '无变更'],
    ['l5_temp.py', 'L5 温度采集 (LHM DLL)', '无变更'],
]
add_table(aux_table[0], aux_table[1:])

doc.add_page_break()

# ============================================================
# 4. V6.2 修复清单
# ============================================================
add_heading_cn('4. V6.2 修复清单', 1)

add_para('本次审查发现并修复了 13 个问题，覆盖 P0（崩溃）、P1（逻辑错误）、P2（代码质量）三个优先级。', bold=True)

add_heading_cn('4.1 P0 — 运行时崩溃', 2)
p0_table = [
    ['#', '文件', '问题', '修复'],
    ['1', 'quadrant_stabilizer.py', 'oscillation_suppress 属性未定义，震荡检测时崩溃', '新增参数 + self 赋值'],
    ['2', 'template_engine.py', 'except: 裸 catch 吞掉所有异常', '改为 except (ValueError, OverflowError)'],
]
add_table(p0_table[0], p0_table[1:])

add_heading_cn('4.2 P1 — 逻辑错误', 2)
p1_table = [
    ['#', '文件', '问题', '修复'],
    ['3', 'ema_filter.py', '默认参数 (0.10, 0.70) 与 V6 训练值 (0.35, 0.60) 不一致', '改为训练最优值'],
    ['4', 'real_collector.py', 'io_congestion 用累计计数器，首帧后趋近 1.0', '改用增量计算'],
    ['5', 'real_collector.py', '_thread_cache 懒初始化脆弱', '移到 __init__ 显式声明'],
    ['7', 'context_pad.py', 'compose_pad() 绕过 BodySenseManager，F/T/C 无记忆', '新增 body 参数，优先使用 EMA 累积值'],
    ['8', 'ode_dynamics.py', 'compute_target() ~80 行与 pad_model 重复', '委托给 pad_mapping.compute_pad_raw()'],
]
add_table(p1_table[0], p1_table[1:])

add_heading_cn('4.3 P2 — 代码质量', 2)
p2_table = [
    ['#', '文件', '问题', '修复'],
    ['9', 'pad_model.py', 'docstring 写 V5.0 但实际是 V4.1', '更正版本标签'],
    ['10', 'pad_mapping.py', 'docstring 写 V5.0', '更正为 V6.0'],
    ['13', 'template_engine.py', 'OutputThrottler 用步数计数而非真实时间', '改用 time.monotonic()'],
    ['—', 'cross_validate.py', 'hysteresis= 参数已删除，TypeError', '改用上下文自适应参数'],
    ['—', 'final_validation.py', '同上', '同上'],
    ['—', 'tune_context_pad.py', '同上', '同上'],
    ['—', 'tune_stabilizer.py', '同上', '同上'],
]
add_table(p2_table[0], p2_table[1:])

add_heading_cn('4.4 管线打通验证', 2)
add_para('BodySense 管线集成测试结果：')
pipe_table = [
    ['指标', 'WITH BodySense', 'WITHOUT BodySense', '差异'],
    ['F (疲劳)', '0.800', '0.433', '+85% (EMA 累积)'],
    ['T (紧绑)', '0.000', '0.092', '窗口紧绑 vs 瞬时'],
    ['C (舒适)', '0.555', '0.719', '资源余量 vs 瞬时'],
]
add_table(pipe_table[0], pipe_table[1:])
add_para('结论：BodySense 的累积疲劳比瞬时值高 85%，长时间高负载的"记忆"能正确传递到 ODE。')

doc.add_page_break()

# ============================================================
# 5. V6.3 新增：ODE-Kalman 融合管线
# ============================================================
add_heading_cn('5. V6.3 新增：ODE-Kalman 融合管线 (signal_pipeline.py)', 1)

add_heading_cn('5.1 设计动机', 2)
add_para(
    'V6.2 的 EMA 滤波器虽然能在平滑性和响应性之间动态平衡，但存在两个根本问题：'
)
add_para(
    '1. 串行延迟叠加：EMA 平滑 → ODE 动力学 → 防闪烁，三层独立处理，每层都有延迟，叠加后总延迟不可控。'
)
add_para(
    '2. 无预测能力：EMA 只能平滑历史数据，无法预测下一步状态。ODE 有动力学模型但无法利用观测修正。'
)
add_para(
    'V6.3 引入 ODE-Kalman 融合管线：将 ODE 动力学模型作为 Kalman 滤波器的预测模型，'
    '观测值作为修正。融合后 ODE 的物理含义（惯性/阻尼）直接进入 Kalman 的预测精度，'
    '预测越准，所需增益越小，延迟越短。'
)

add_heading_cn('5.2 三级信号链', 2)
add_para('V6.3 信号管线采用三级信号处理链：', bold=True)
signal_chain = [
    ['层级', '模块', '功能', '默认配置'],
    ['L1', 'MedianSpikeFilter', '3点中值滤波，去除尖峰噪声', 'enable=False（本地数据干净）'],
    ['L2', 'ODEKalmanPAD', 'ODE预测 + Kalman修正，自适应q', 'q_base=0.001, r=0.02'],
    ['L3', 'DualHysteresis', '象限滞回 + 情绪标签滞回', 'margin=0.03, min_hold=3'],
]
add_table(signal_chain[0], signal_chain[1:])

add_heading_cn('5.3 ODE-Kalman 融合原理', 2)
add_para('核心方程：', bold=True)
add_para('预测步（ODE）：x_hat_minus = x_hat_prev + dt * (-k_spring * (x_hat_prev - target) - gamma * v_prev)')
add_para('修正步（Kalman）：x_hat = x_hat_minus + K * (z - x_hat_minus)')
add_para('增益自适应：K = P_minus / (P_minus + r)，P_minus = P_prev + q')
add_para('自适应 q（NIS）：若 NIS > 阈值，则 q = q_base * (1 + NIS_excess)，观测异常时提高预测权重')

ode_kalman_params = [
    ['参数', '训练最优值', '含义'],
    ['gamma', '0.2', 'ODE 阻尼系数，控制震荡衰减速度'],
    ['k_spring', '0.05', 'ODE 弹簧系数，控制向目标回归速度'],
    ['q_base', '0.001', 'Kalman 基础过程噪声，反映模型预测精度'],
    ['r', '0.02', 'Kalman 观测噪声，反映传感器噪声水平'],
    ['quadrant_margin', '0.03', '象限滞回死区宽度'],
    ['emotion_min_hold', '3', '情绪标签最小保持步数'],
    ['emotion_intensity_margin', '0.1', '情绪强度切换裕度'],
]
add_table(ode_kalman_params[0], ode_kalman_params[1:])

add_heading_cn('5.4 残差监控机制', 2)
add_para(
    '信号管线内置残差监控：每步计算观测值 PAD 与 Kalman 估计值之间的残差，'
    '在滚动窗口（100步）内统计方差。残差方差反映 Kalman 模型的预测精度。'
)
residual_thresholds = [
    ['阈值', '含义', '触发动作'],
    ['< 0.03', '模型匹配良好', '无需操作'],
    ['0.03 ~ 0.06', '模型精度下降', '关注（建议记录）'],
    ['> 0.06', '模型失配', '警告（建议上调 q 或 r）'],
]
add_table(residual_thresholds[0], residual_thresholds[1:])

doc.add_page_break()

# ============================================================
# 6. V6.3 新增：真机压力测试
# ============================================================
add_heading_cn('6. V6.3 新增：真机压力测试', 1)

add_heading_cn('6.1 测试目标', 2)
add_para(
    'V6.2 的测试均在离线/模拟环境下完成。V6.3 需要在真实机器上验证：'
    '（1）CPU 高负载下参数是否稳定；（2）残差方差是否随负载升高而增大；'
    '（3）gradual overload 场景下象限切换是否平滑。'
)

add_heading_cn('6.2 测试设计', 2)
add_para('demo_stress.py 使用 multiprocessing 启动子进程打满 CPU：', bold=True)
stress_design = [
    ['阶段', '进程数', '步数', '目标 CPU', '实际 CPU 范围'],
    ['idle', '0', '10', '10~20%', '11~19%'],
    ['stress_1p', '1', '10', '20~30%', '17~28%'],
    ['stress_2p', '2', '10', '30~40%', '20~39%'],
    ['stress_4p', '4', '15', '45~55%', '43~53%'],
    ['stress_8p', '8', '15', '75~95%', '75~92%'],
    ['recovery', '0', '25', '释放后恢复', '84%下降中'],
]
add_table(stress_design[0], stress_design[1:])

add_heading_cn('6.3 测试结果', 2)
add_para('关键发现：', bold=True)
stress_results = [
    ['指标', 'idle', '1p', '2p', '4p', '8p', '判定'],
    ['CPU avg', '14.6%', '22.3%', '30.7%', '47.0%', '80.8%', '平滑递增'],
    ['CPU max', '19.4%', '28.4%', '39.4%', '52.8%', '92.5%', '推到 90%+'],
    ['残差P方差', '0.0000', '0.0000', '0.0000', '0.0000', '0.0000', '始终保持 0'],
    ['残差A方差', '0.0048', '0.0022', '0.0027', '0.0017', '0.0065', '远低于 r=0.02'],
    ['象限切换', '—', '2次', '1次', '0次', '0次', '平滑过渡'],
]
add_table(stress_results[0], stress_results[1:])
add_para(
    '结论：q=0.001 / r=0.02 在 CPU 92.5% 时残差方差仅 0.0065，远低于 r=0.02 的观测噪声水平。'
    '参数在真机高负载下完全稳定，无需调整。'
)

add_heading_cn('6.4 离线测试回顾', 2)
add_para('V6.2 离线测试已覆盖 5 个场景（idle / stress_1t / stress_4t / stress_8t / recovery），共 142 行数据，4 种情绪跨象限分布。')
add_para('V6.3 真机测试补充了 gradual_overload 逐级加码场景，验证了 multiprocessing 压力注入的有效性。')

doc.add_page_break()

# ============================================================
# 7. V6.3 新增：7 天连续采集
# ============================================================
add_heading_cn('7. V6.3 新增：7 天连续采集与长期验证', 1)

add_heading_cn('7.1 采集方案', 2)
add_para(
    'V6.2 仅有 30 分钟低负载采集数据。V6.3 部署了 7 天连续采集方案，覆盖不同负载模式。'
)
collect_plan = [
    ['采集时段', '时间', '时长', '覆盖场景', '预期数据量'],
    ['工作日高峰', '每天 10:00-11:00', '1 小时', '高负载、多任务', '~7 x 3600 行'],
    ['夜间空闲', '每天 02:00-03:00', '1 小时', '低负载、后台任务', '~7 x 3600 行'],
    ['合计', '—', '~14 小时', '7 天 x 2 时段', '~50,000 行'],
]
add_table(collect_plan[0], collect_plan[1:])

add_heading_cn('7.2 双重用途', 2)
add_para('1. 参数验证：用 14 小时连续数据重新估计 q/r，验证训练最优参数在长周期下的稳定性。')
add_para('2. LLM 微调数据源：时序链数据（原始指标→PAD→情绪→表达）可作为 LLM 微调的训练数据，使 LLM 理解计算机情绪状态。')

add_heading_cn('7.3 LiveRecorder 模块', 2)
add_para(
    'LiveRecorder（live_recorder.py）将实时采集数据写入 CSV 表格，包含 24 列：'
    '时间戳、步数、阶段标签、CPU/内存/Swap、原始PAD、滤波后PAD、残差P/A/D、'
    'Kalman增益、象限、情绪标签、体感维度、输出文本。'
)
add_para('每次采集完成后自动汇报：总步数、象限分布、残差方差、闪烁率。')

doc.add_page_break()

# ============================================================
# 8. 测试验证
# ============================================================
add_heading_cn('8. 测试验证', 1)

add_heading_cn('8.1 模块导入测试', 2)
add_para('15/15 核心模块全部导入成功，0 失败。')

add_heading_cn('8.2 QuadrantStabilizer 测试', 2)
add_para('4 个场景全部通过：稳定信号不切换、噪声振荡不切换、真实变化正确响应、回来需惯性窗口确认。')

add_heading_cn('8.3 ODE 动力学测试', 2)
add_para('空闲→突发异常→恢复：P 骤降时 A 飙升（耦合效应），恢复后 P 回升。')

add_heading_cn('8.4 Plutchik 情感轮测试', 2)
add_para('8 个测试用例全部通过：空闲→喜悦，严重过载→愤怒，突发异常→恐惧，恢复平静→信任。')

add_heading_cn('8.5 BodySense 体感测试', 2)
add_para('5 个场景通过：空闲、L2 连接风暴、L3 IO 阻塞、L5 物理过热、全层并发（噩梦场景）。')

doc.add_page_break()

# ============================================================
# 6. 关键技术指标
# ============================================================
add_heading_cn('9. 关键技术指标', 1)

perf_table = [
    ['指标', '数值', '说明'],
    ['原始指标', '58', '五层感官采集'],
    ['派生指标', '29', '实时计算'],
    ['总指标数', '87', '58 + 29'],
    ['体感维度', '3', '疲劳/紧绑/舒适'],
    ['语义信号', '4', 'error/load/latency/health'],
    ['PAD 维度', '4', 'P/A/D/V'],
    ['情感维度', '7', 'P/A/D/V/F/T/C'],
    ['Plutchik 状态', '24', '8 基本情绪 x 3 强度'],
    ['复合情绪', '8', '相邻二元组合'],
    ['运维意图簇', '5', '降维映射'],
    ['Kalman 参数', 'q=0.001, r=0.02', '训练最优，真机验证通过'],
    ['CPU 极限测试', '92.5%', '残差方差仅 0.0065'],
    ['LLM 成本', '0 元', 'Phase 1 纯规则+模板'],
    ['单步延迟', '<1ms', '纯 Python'],
]
add_table(perf_table[0], perf_table[1:])

doc.add_page_break()

# ============================================================
# 7. 多AI协作优化记录
# ============================================================
add_heading_cn('10. 多AI协作优化记录', 1)
add_para(
    '本项目所有重大设计决策均经过DeepSeek、Qwen、GLM三个AI模型的独立评审，'
    '再由主模型进行交叉验证和最终整合。'
)

collab_table = [
    ['优化项', 'DeepSeek', 'Qwen', 'GLM', '最终方案'],
    ['PAD映射', '分段Sigmoid', '乘法D值', '健康系数headroom', '三者整合'],
    ['ODE动力学', '惯性+爆发+衰减', '外部驱动F/T/C', '自适应衰减率', '三者整合'],
    ['A值坍塌修复', '自适应衰减率', '自适应ODE+断崖检测', '断崖+软边界+惯性反转', '三重机制整合'],
    ['上下文感知', '语义信号层', '上下文PAD映射', '体感累积', 'V6 四层管线'],
    ['防闪烁', '维度死区', '象限滞回', '惯性窗口+震荡抑制', '四层防抖'],
    ['EMA训练', 'Google真实数据', '网格搜索', '多目标评分', '720组参数最优'],
    ['Kalman管线', 'ODE-Kalman融合', '自适应q(NIS)', '残差监控', '融合管线+训练验证'],
]
add_table(collab_table[0], collab_table[1:])

doc.add_page_break()

# ============================================================
# 8. 项目文件清单
# ============================================================
add_heading_cn('11. 项目文件清单', 1)

files_table = [
    ['文件名', '功能描述', 'V6.3 状态'],
    ['real_collector.py', '五层感官采集器（58原始+29派生指标）', 'io_congestion 修复'],
    ['body_sense.py', '体感维度（疲劳/紧绑/舒适 EMA 追踪）', '正常'],
    ['semantic_signals.py', '语义信号层（4 正交信号 + 上下文标签）', '正常'],
    ['context_pad.py', '上下文感知 PAD 映射', 'body 参数新增'],
    ['pad_mapping.py', 'PAD 映射共享模块', '版本标签更新'],
    ['pad_model.py', 'PAD 状态、6 象限分类、健康分', '版本标签更新'],
    ['ode_dynamics.py', 'ODE 动力系统（惯性/爆发/衰减）', 'compute_target 去重'],
    ['signal_pipeline.py', 'ODE-Kalman 融合管线（V6.3 新增）', '真机验证通过'],
    ['ema_filter.py', '自适应 EMA + 情绪惯性', '默认参数修正'],
    ['quadrant_stabilizer.py', '防闪烁象限控制器（四层防抖）', 'oscillation_suppress 修复'],
    ['plutchik.py', 'Plutchik 情感轮（8x3+8复合）', '正常'],
    ['template_engine.py', '模板引擎（6状态x5模板）', 'except 修复 + 时间修复'],
    ['habituation.py', 'Weber-Fechner 防疲劳表达', '正常'],
    ['dirty_buffer.py', '脏数据缓冲（空值/断崖/乱序）', '正常'],
    ['live_recorder.py', '实时数据采集录制（V6.3 新增）', '24列CSV输出'],
    ['l4_metrics.py', 'L4 Prometheus 采集', '正常'],
    ['l4_proxy.py', 'L4 HTTP 代理', '正常'],
    ['l5_temp.py', 'L5 温度采集 (LHM DLL)', '正常'],
    ['train_ema.py', 'EMA 参数网格搜索训练', '版本标签更新'],
    ['train_pipeline.py', '信号管线参数训练（V6.3 新增）', '720组网格搜索'],
    ['main_live.py', 'V6 实时引擎（完整四层管线）', 'LiveRecorder 集成'],
    ['demo_stress.py', 'V6.3 真机压力测试', 'mutiprocessing 压力注入'],
    ['main.py', 'V3/V4 遗留入口', '标记为遗留'],
    ['cross_validate.py', '交叉验证', 'hysteresis 修复'],
    ['final_validation.py', '最终验证', 'hysteresis 修复'],
    ['tune_stabilizer.py', 'Stabilizer 参数调优', 'hysteresis 修复'],
    ['tune_context_pad.py', 'Context PAD 参数调优', 'hysteresis 修复'],
]
add_table(files_table[0], files_table[1:])

doc.add_page_break()

# ============================================================
# 9. 结论与展望
# ============================================================
add_heading_cn('12. 结论与展望', 1)

add_heading_cn('12.1 V6 项目成果', 2)
achievements = [
    '构建了完整的四层浓缩管线：87指标 -> 3体感 -> 4语义 -> PAD -> 7维情感',
    '实现上下文感知 PAD 映射：同一指标在不同语境下产生不同情绪',
    'BodySense 体感累积系统：疲劳/紧绑/舒适具有 EMA 记忆',
    'ODE 动力系统：情绪具有惯性、爆发、衰减特性',
    'V6.3 ODE-Kalman 融合管线：ODE 作为 Kalman 预测模型，单层延迟',
    'V6.3 真机压力测试：CPU 92.5% 时残差方差仅 0.0065，参数稳定',
    'V6.3 7 天连续采集：覆盖高峰/夜间，双重用途（参数验证+LLM数据）',
    'V6.2 审查修复：13 个问题（2 P0 + 5 P1 + 6 P2），管线完全打通',
    '15/15 核心模块导入通过，全部自测试通过',
]
for a in achievements:
    doc.add_paragraph(a, style='List Bullet')

add_heading_cn('12.2 后续展望', 2)
future = [
    'cross_validate.py / final_validation.py 数据结构适配',
    'main.py 迁移到 V6 BodySense 管线',
    'Phase 2：引入 LLM 生成动态情绪化表达',
    'Phase 3：接入 Mate-Engine 项目验证工程实用性',
    'Phase 4：基于历史数据建立情绪基线，实现异常预测',
    'Phase 5：7 天数据用于 LLM 微调，使 LLM 理解计算机情绪状态',
]
for f in future:
    doc.add_paragraph(f, style='List Bullet')

# ============================================================
# 附录
# ============================================================
doc.add_page_break()
add_heading_cn('附录A：项目开发时间线', 1)
timeline = [
    ['阶段', '时间', '里程碑'],
    ['V1-V3', '第1-2天', '基础PAD映射+Plutchik+ODE+EMA+模板'],
    ['V4-V5', '第2-3天', 'A值坍塌修复+黑盒测试+压力测试+EMA训练'],
    ['V6.0', '第3天', '上下文感知架构：semantic_signals + context_pad + body_sense'],
    ['V6.1', '第4天', 'pad_mapping 恢复+文档更新+已知缺陷记录'],
    ['V6.2', '第4天', '代码审查修复：13个问题，管线完全打通，报告更新'],
    ['V6.3', '第5天', 'ODE-Kalman融合+真机压力测试(CPU92.5%)+7天采集部署'],
]
add_table(timeline[0], timeline[1:])

# === 保存 ===
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), '计算机情绪引擎项目报告_V6.3.docx')
doc.save(output_path)
print(f'报告已生成: {output_path}')
print(f'文件大小: {os.path.getsize(output_path) / 1024:.1f} KB')
